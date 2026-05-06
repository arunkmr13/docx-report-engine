from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


DEFAULT_SECTIONS = {
    "Executive Summary": [
        {"type": "paragraph", "text": "Overview of key business outcomes, major achievements, and strategic direction."}
    ],
    "Market Analysis": [
        {"type": "paragraph", "text": "Insights into current market trends, competitor positioning, and growth opportunities."}
    ],
    "Department Highlights": [
        {"type": "paragraph", "text": "Performance review of each department with key metrics and accomplishments."}
    ],
    "Financial Overview": [
        {"type": "paragraph", "text": "Breakdown of revenue, expenses, and profitability across business units."}
    ],
    "Future Strategy": [
        {"type": "paragraph", "text": "Planned initiatives, expansion goals, and strategic priorities for upcoming quarters."}
    ]
}


def generate_content(doc, sections: dict = None):
    if sections is None:
        sections = DEFAULT_SECTIONS

    for idx, (section_name, blocks) in enumerate(sections.items()):
        if idx != 0:
            doc.add_page_break()

        # Section heading
        doc.add_heading(f"Section {idx + 1}: {section_name}", level=1)

        # Render each block
        for block in blocks:
            block_type = block.get("type")

            if block_type == "paragraph":
                doc.add_paragraph(block.get("text", ""))

            elif block_type == "bullet_list":
                for item in block.get("items", []):
                    doc.add_paragraph(item, style="List Bullet")

            elif block_type == "numbered_list":
                for item in block.get("items", []):
                    doc.add_paragraph(item, style="List Number")

            elif block_type == "heading":
                level = block.get("level", 2)
                doc.add_heading(block.get("text", ""), level=level)

            elif block_type == "table":
                headers = block.get("headers", [])
                rows = block.get("rows", [])

                if not headers and not rows:
                    continue

                col_count = len(headers) if headers else len(rows[0])
                table = doc.add_table(
                    rows=1 + len(rows),
                    cols=col_count
                )
                table.style = "Table Grid"

                # Header row
                if headers:
                    header_row = table.rows[0].cells
                    for i, h in enumerate(headers):
                        header_row[i].text = h
                        for para in header_row[i].paragraphs:
                            for run in para.runs:
                                run.bold = True

                # Data rows
                for r_idx, row_data in enumerate(rows):
                    row_cells = table.rows[r_idx + 1].cells
                    for c_idx, cell_val in enumerate(row_data):
                        row_cells[c_idx].text = str(cell_val)

                doc.add_paragraph()  # spacing after table

            elif block_type == "image":
                path = block.get("path", "")
                width = block.get("width", 4.0)
                align = block.get("align", "left")
                try:
                    para = doc.add_paragraph()
                    if align == "center":
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif align == "right":
                        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    run = para.add_run()
                    run.add_picture(path, width=Inches(width))
                except Exception:
                    doc.add_paragraph(f"[Image not found: {path}]")

            elif block_type == "page_break":
                doc.add_page_break()

            elif block_type == "divider":
                para = doc.add_paragraph()
                pPr = para._p.get_or_add_pPr()
                pBdr = doc.element.makeelement(
                    qn('w:pBdr'), {}
                )
                bottom = doc.element.makeelement(
                    qn('w:bottom'), {
                        qn('w:val'): 'single',
                        qn('w:sz'): '6',
                        qn('w:space'): '1',
                        qn('w:color'): 'AAAAAA'
                    }
                )
                pBdr.append(bottom)
                pPr.append(pBdr)

            else:
                # Unknown block type — skip gracefully
                pass