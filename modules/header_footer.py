from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _clear_hf(hf_element):
    """Nuclear clear — removes every child XML node."""
    for child in list(hf_element):
        hf_element.remove(child)
    hf_element.text = None
    hf_element.tail = None


def _remove_first_page_references(section):
    """Remove 'first' and 'even' type header/footer references from sectPr."""
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    sectPr = section._sectPr
    for ref in list(sectPr):
        tag = ref.tag.split('}')[1] if '}' in ref.tag else ref.tag
        if tag in ('headerReference', 'footerReference'):
            ref_type = ref.get(f'{{{ns}}}type')
            if ref_type in ('first', 'even'):
                sectPr.remove(ref)


def _disable_even_odd_headers(doc):
    """Disable evenAndOddHeaders setting in document settings."""
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    settings = doc.settings.element
    even_odd = settings.find(f'{{{ns}}}evenAndOddHeaders')
    if even_odd is not None:
        settings.remove(even_odd)


def _remove_table_borders(table):
    """Remove all borders from a header/footer table."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def apply_header_footer(doc, title, quarter, company_name, prepared_by, header_font_size=10, footer_font_size=10, logo_path=None, logo_size=0.5, logo_position="left", page_label="Page no: ", extra_fields=None):

    _disable_even_odd_headers(doc)

    for section in doc.sections:

        section.different_first_page_header_footer = False
        _remove_first_page_references(section)
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False

        usable_width = section.page_width - section.left_margin - section.right_margin

        # ---------------- HEADER ----------------
        header = section.header
        _clear_hf(header._element)

        if logo_path:
            table = header.add_table(rows=1, cols=3, width=usable_width)
            _remove_table_borders(table)
            cell_a = table.rows[0].cells[0]
            cell_b = table.rows[0].cells[1]
            cell_c = table.rows[0].cells[2]

            if logo_position == "left":
                cell_a.width = int(usable_width * 0.12)
                cell_b.width = int(usable_width * 0.33)
                cell_c.width = int(usable_width * 0.55)
                logo_cell = cell_a
                mid_cell  = cell_b
                text_cell = cell_c
                mid_cell.paragraphs[0].add_run(company_name)
                mid_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                text_cell.paragraphs[0].add_run(f"{title} | {quarter}")
                text_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            elif logo_position == "center":
                cell_a.width = int(usable_width * 0.35)
                cell_b.width = int(usable_width * 0.12)
                cell_c.width = int(usable_width * 0.53)
                logo_cell = cell_b
                mid_cell  = cell_a
                text_cell = cell_c
                mid_cell.paragraphs[0].add_run(company_name)
                mid_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                text_cell.paragraphs[0].add_run(f"{title} | {quarter}")
                text_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            elif logo_position == "right":
                cell_a.width = int(usable_width * 0.35)
                cell_b.width = int(usable_width * 0.53)
                cell_c.width = int(usable_width * 0.12)
                logo_cell = cell_c
                mid_cell  = cell_a
                text_cell = cell_b
                mid_cell.paragraphs[0].add_run(company_name)
                mid_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                text_cell.paragraphs[0].add_run(f"{title} | {quarter}")
                text_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            else:
                raise ValueError(f"Invalid logo_position '{logo_position}'. Use 'left', 'center', or 'right'.")

            logo_para = logo_cell.paragraphs[0]
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = logo_para.add_run()
            run.add_picture(logo_path, width=Inches(logo_size))

            for cell in [mid_cell, text_cell]:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(header_font_size)

        else:
            table = header.add_table(rows=1, cols=2, width=usable_width)
            _remove_table_borders(table)
            left_cell  = table.rows[0].cells[0]
            right_cell = table.rows[0].cells[1]
            left_cell.width  = int(usable_width * 0.35)
            right_cell.width = int(usable_width * 0.65)
            left_cell.paragraphs[0].add_run(company_name)
            left_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            right_cell.paragraphs[0].add_run(f"{title} | {quarter}")
            right_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for cell in [left_cell, right_cell]:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(header_font_size)

        # ---------------- FOOTER ----------------
        footer = section.footer
        _clear_hf(footer._element)

        extras = extra_fields or {}
        footer_rows = 2 if extras else 1

        table = footer.add_table(rows=footer_rows, cols=2, width=usable_width)
        _remove_table_borders(table)

        # Row 1 — page number (left) + prepared by (right)
        left_cell  = table.rows[0].cells[0]
        right_cell = table.rows[0].cells[1]
        left_cell.width  = int(usable_width * 0.35)
        right_cell.width = int(usable_width * 0.65)

        # Left cell — page label + page number field
        para = left_cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(page_label)
        run.font.size = Pt(footer_font_size)

        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

        # Right cell — prepared by
        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right_run = right_para.add_run(prepared_by)
        right_run.font.size = Pt(footer_font_size)

        # Row 2 — extra fields
        if extras:
            extra_left  = table.rows[1].cells[0]
            extra_right = table.rows[1].cells[1]
            extra_left.width  = int(usable_width * 0.50)
            extra_right.width = int(usable_width * 0.50)

            items = list(extras.items())
            mid = len(items) // 2 + len(items) % 2
            left_items  = items[:mid]
            right_items = items[mid:]

            left_text  = "  |  ".join(f"{k}: {v}" for k, v in left_items)
            right_text = "  |  ".join(f"{k}: {v}" for k, v in right_items)

            left_extra_para = extra_left.paragraphs[0]
            left_extra_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            left_extra_para.add_run(left_text).font.size = Pt(footer_font_size - 1)

            if right_text:
                right_extra_para = extra_right.paragraphs[0]
                right_extra_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                right_extra_para.add_run(right_text).font.size = Pt(footer_font_size - 1)