from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from lxml import etree


# Register VML namespace
VMAP = 'urn:schemas-microsoft-com:vml'
VMAP_O = 'urn:schemas-microsoft-com:office:office'
VMAP_W10 = 'urn:schemas-microsoft-com:office:word'

etree.register_namespace('v', VMAP)
etree.register_namespace('o', VMAP_O)
etree.register_namespace('w10', VMAP_W10)


def apply_watermark(doc, text="SAMPLE"):
    for section in doc.sections:
        header = section.header

        if not header.paragraphs:
            header.add_paragraph()

        para = header.paragraphs[0]

        pict = OxmlElement('w:pict')

        shape = etree.SubElement(pict, '{%s}shape' % VMAP)
        shape.set('id', 'watermark')
        shape.set('type', '#_x0000_t136')
        shape.set('style',
            'position:absolute;'
            'margin-left:0;'
            'margin-top:0;'
            'width:527pt;'
            'height:144pt;'
            'z-index:-251654144;'
            'mso-position-horizontal:center;'
            'mso-position-horizontal-relative:margin;'
            'mso-position-vertical:center;'
            'mso-position-vertical-relative:margin'
        )
        shape.set('fillcolor', '#d0d0d0')
        shape.set('stroked', 'f')

        fill = etree.SubElement(shape, '{%s}fill' % VMAP)
        fill.set('on', 't')
        fill.set('focussize', '0,0')

        textpath = etree.SubElement(shape, '{%s}textpath' % VMAP)
        textpath.set('style', 'font-family:"Calibri";font-size:1pt')
        textpath.set('string', text)

        para._p.append(pict)